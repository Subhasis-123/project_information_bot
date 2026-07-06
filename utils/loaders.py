import os
import pandas as pd

from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain_core.documents import Document


class DocumentLoader:
    """
    Enterprise Document Loader

    Responsibilities:
    -----------------
    1. Load supported documents
    2. Convert them into LangChain Documents
    3. Attach metadata
    """

    def __init__(self):

        self.supported_files = {

            ".pdf": self.load_pdf,

            ".docx": self.load_docx,

            ".txt": self.load_txt,

            ".csv": self.load_csv,

            ".md": self.load_md

        }

    # ======================================================
    # PDF
    # ======================================================

    def load_pdf(self, file_path):

        documents = []

        reader = PdfReader(file_path)

        for page_no, page in enumerate(reader.pages):

            text = page.extract_text()

            if text:

                documents.append(

                    Document(

                        page_content=text,

                        metadata={

                            "source": os.path.basename(file_path),

                            "file_type": "PDF",

                            "page": page_no + 1

                        }

                    )

                )

        return documents

    # ======================================================
    # DOCX
    # ======================================================

    def load_docx(self, file_path):

        documents = []

        doc = DocxDocument(file_path)

        text = "\n".join(

            para.text

            for para in doc.paragraphs

            if para.text.strip()

        )

        documents.append(

            Document(

                page_content=text,

                metadata={

                    "source": os.path.basename(file_path),

                    "file_type": "DOCX"

                }

            )

        )

        return documents

    # ======================================================
    # TXT
    # ======================================================

    def load_txt(self, file_path):

        documents = []

        with open(file_path, "r", encoding="utf-8") as f:

            text = f.read()

        documents.append(

            Document(

                page_content=text,

                metadata={

                    "source": os.path.basename(file_path),

                    "file_type": "TXT"

                }

            )

        )

        return documents

    # ======================================================
    # CSV
    # ======================================================

    def load_csv(self, file_path):

        documents = []

        df = pd.read_csv(file_path)

        for _, row in df.iterrows():

            text = "\n".join(

                f"{col}: {row[col]}"

                for col in df.columns

            )

            documents.append(

                Document(

                    page_content=text,

                    metadata={

                        "source": os.path.basename(file_path),

                        "file_type": "CSV"

                    }

                )

            )

        return documents

    # ======================================================
    # Markdown
    # ======================================================

    def load_md(self, file_path):

        documents = []

        with open(file_path, "r", encoding="utf-8") as f:

            text = f.read()

        documents.append(

            Document(

                page_content=text,

                metadata={

                    "source": os.path.basename(file_path),

                    "file_type": "Markdown"

                }

            )

        )

        return documents

    # ======================================================
    # Main Loader
    # ======================================================

    def load_documents(self, folder):

        documents = []

        if not os.path.exists(folder):

            raise FileNotFoundError(

                f"Knowledge folder '{folder}' not found."

            )

        for file in os.listdir(folder):

            file_path = os.path.join(folder, file)

            if not os.path.isfile(file_path):

                continue

            extension = os.path.splitext(file)[1].lower()

            loader = self.supported_files.get(extension)

            if loader:

                try:

                    documents.extend(

                        loader(file_path)

                    )

                except Exception as e:

                    print(

                        f"Error loading {file}: {e}"

                    )

        return documents


# ======================================================
# Singleton Instance
# ======================================================

document_loader = DocumentLoader()